from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_PATH = PROJECT_ROOT / "Documento_Sprints_Chatbot_PUCESI.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_heading(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, style: str | None = None, bold_prefix: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)
    return paragraph


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        add_para(doc, item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        add_para(doc, item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr_cells[i], header, bold=True)
        set_cell_shading(hdr_cells[i], "D9EAF7")
        hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            hdr_cells[i].width = Inches(widths[i])

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths:
                cells[i].width = Inches(widths[i])

    doc.add_paragraph()
    return table


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(31, 78, 121)

    for name, size, color in [
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 11.5, RGBColor(68, 68, 68)),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    for list_style_name in ["List Bullet", "List Number"]:
        list_style = styles[list_style_name]
        list_style.font.name = "Calibri"
        list_style.font.size = Pt(10.5)
        list_style.paragraph_format.space_after = Pt(3)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Documento de Desarrollo por Sprints\n")
    subtitle = title.add_run("Chatbot Académico Inteligente PUCE-SI")
    subtitle.font.size = Pt(18)
    subtitle.font.color.rgb = RGBColor(46, 116, 181)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Análisis funcional y técnico del proyecto por iteraciones Scrum").italic = True

    doc.add_paragraph()
    add_table(
        doc,
        ["Campo", "Detalle"],
        [
            ["Proyecto", "Prototipo de chatbot académico inteligente para la Pontificia Universidad Católica del Ecuador Sede Ibarra."],
            ["Enfoque", "Consultas académicas personalizadas, consultas institucionales mediante RAG, interacción por voz, interfaz conversacional, seguridad y auditoría."],
            ["Alcance", "Backend Django REST Framework, frontend React/Vite/TypeScript, base académica, RAG institucional, OpenAI para lenguaje/voz, trazabilidad de auditoría y controles de seguridad."],
            ["Fecha de elaboración", "Junio de 2026"],
        ],
        widths=[1.7, 5.2],
    )
    doc.add_page_break()


def add_intro(doc: Document) -> None:
    add_heading(doc, "1. Propósito del documento", 1)
    add_para(
        doc,
        "Este documento describe de forma detallada el contenido de los cinco sprints del proyecto. "
        "La explicación se basa en la arquitectura real implementada: aplicaciones del backend, modelos de datos, endpoints, servicios de inteligencia artificial, motor RAG, componentes del frontend, controles de seguridad y mecanismos de auditoría.",
    )
    add_para(
        doc,
        "El objetivo es dejar una trazabilidad clara entre cada sprint y las funcionalidades desarrolladas. "
        "Por ello no se limita a describir historias de usuario; también explica cómo cada historia se materializa en módulos, servicios, flujos de datos y pantallas del prototipo.",
    )

    add_heading(doc, "2. Vista general de la arquitectura", 1)
    add_para(
        doc,
        "El proyecto está organizado en dos bloques principales. El backend expone servicios REST con Django REST Framework y contiene las aplicaciones de autenticación, estudiantes, auditoría, chatbot, RAG y voz. "
        "El frontend está construido con React, Vite y TypeScript, y consume dichos servicios mediante un cliente HTTP con manejo de tokens JWT y refresco automático de sesión.",
    )
    add_table(
        doc,
        ["Capa", "Componentes principales", "Responsabilidad"],
        [
            [
                "Frontend",
                "React, Vite, TypeScript, Zustand, Axios, rutas públicas y protegidas, ChatWindow, LandingPage, LoginPage, DashboardPage, TeacherDashboardPage.",
                "Presentar la experiencia conversacional y académica, gestionar sesión del usuario, mostrar información académica, permitir carga de actividades y habilitar interacción textual y por voz.",
            ],
            [
                "Backend API",
                "Django REST Framework, autenticación JWT, endpoints de estudiantes, endpoints de chatbot, voz, RAG y auditoría.",
                "Centralizar la lógica de negocio, validar permisos, consultar la base académica, invocar IA, recuperar información institucional y registrar eventos críticos.",
            ],
            [
                "Base académica",
                "Student, Professor, Classroom, Subject, Schedule, Enrollment, Grade, Activity, ActivitySubmission, Notification, ConversationHistory.",
                "Representar perfiles, carreras, materias, horarios, calificaciones, actividades, entregas, notificaciones e historial conversacional autenticado.",
            ],
            [
                "RAG institucional",
                "Documentos oficiales, indexación FAISS, embeddings de OpenAI, búsqueda semántica, respaldo léxico, jerarquía de fuentes.",
                "Responder consultas institucionales usando fuentes validadas de PUCE-SI y documentos internos actualizados.",
            ],
            [
                "IA conversacional y voz",
                "OpenAI Chat Completions, Whisper/STT, TTS, herramientas de fecha/hora/clima, contexto académico y contexto RAG.",
                "Generar respuestas en español, personalizadas cuando el usuario está autenticado y con soporte de entrada/salida por voz.",
            ],
            [
                "Seguridad y auditoría",
                "JWT, OTP, Argon2, bloqueo por intentos fallidos, throttling, CORS/CSRF, límites de carga, auditoría y endpoints administrativos.",
                "Proteger identidad, datos académicos, archivos, servicios de IA y evidencias de uso del sistema.",
            ],
        ],
        widths=[1.25, 2.55, 3.1],
    )

    add_heading(doc, "3. Resumen de sprints", 1)
    add_table(
        doc,
        ["Sprint", "Nombre del sprint", "Resultado esperado"],
        [
            ["Sprint 1", "Consultas académicas personalizadas", "El estudiante autenticado consulta datos propios como perfil, notas, horarios, estadísticas, notificaciones y actividades, y el chatbot usa ese contexto para responder de forma personalizada."],
            ["Sprint 2", "Consultas institucionales", "El chatbot responde información general de PUCE-SI mediante una arquitectura RAG basada en documentos validados, fuentes institucionales y respaldo de búsqueda léxica."],
            ["Sprint 3", "Interacción conversacional por voz y respuesta hablada", "El sistema permite enviar preguntas por audio, transcribirlas, responderlas con el chatbot y devolver audio hablado al usuario."],
            ["Sprint 4", "Interfaz conversacional", "La aplicación ofrece una experiencia de chat usable, integrada con páginas públicas, dashboard de estudiante, panel docente, login con OTP y manejo de estado de sesión."],
            ["Sprint 5", "Seguridad y auditoría", "El sistema aplica autenticación robusta, roles, protección de datos, limitación de abuso, registro de eventos y consulta administrativa de logs."],
        ],
        widths=[0.9, 2.15, 3.85],
    )


def add_sprint_1(doc: Document) -> None:
    add_heading(doc, "4. Sprint 1: Consultas académicas personalizadas", 1)
    add_para(
        doc,
        "Este sprint construye la base académica personalizada del chatbot. Su finalidad es que el estudiante no reciba respuestas genéricas cuando pregunta por su vida académica, sino respuestas conectadas con su perfil, carrera, semestre, materias, horarios, calificaciones, actividades y notificaciones. "
        "La personalización se activa cuando el usuario inicia sesión correctamente y el backend puede asociar la conversación con un registro Student.",
    )

    add_heading(doc, "4.1 Objetivo funcional", 2)
    add_para(
        doc,
        "El objetivo es permitir consultas académicas privadas en lenguaje natural. Por ejemplo: notas actuales, horario de clases, materias inscritas, profesores asignados, actividades pendientes, entregas realizadas, estadísticas académicas y notificaciones. "
        "La respuesta debe basarse en datos del estudiante autenticado y no debe exponer información de otros usuarios.",
    )

    add_heading(doc, "4.2 Modelo académico implementado", 2)
    add_table(
        doc,
        ["Entidad", "Contenido en el proyecto", "Uso dentro del sprint"],
        [
            [
                "Student",
                "Perfil académico del estudiante: código estudiantil, carrera, semestre actual, fecha de inicio, foto, nombre completo y correo.",
                "Permite identificar al usuario autenticado y enriquecer las respuestas del chatbot con datos personales controlados por sesión.",
            ],
            [
                "Professor",
                "Información del docente: nombres, correo institucional, especialización, estado activo y relación opcional con usuario del sistema.",
                "Se usa para mostrar quién dicta una materia y para permitir consultas del estudiante sobre profesores y horarios.",
            ],
            [
                "Classroom",
                "Código del aula, edificio, piso, capacidad y marca de laboratorio.",
                "Permite que el horario no solo diga materia y hora, sino también ubicación física dentro del campus.",
            ],
            [
                "Subject",
                "Código, nombre, créditos, semestre, carrera y estado activo.",
                "Estructura las materias por carrera y por semestre académico.",
            ],
            [
                "Schedule",
                "Materia, profesor, aula, día de semana, hora de inicio, hora de fin y periodo académico.",
                "Sostiene las consultas de horario y permite que el chatbot responda preguntas como clases de hoy, ubicación o docente asignado.",
            ],
            [
                "Enrollment",
                "Relación entre estudiante y horario para un periodo académico, con estado activo, retirado o completado.",
                "Determina qué materias pertenecen realmente al estudiante en un periodo.",
            ],
            [
                "Grade",
                "Parciales, examen final, nota final sobre 200, asistencia, periodo y estado de aprobación.",
                "Permite consultas de notas, rendimiento, promedios y estado académico.",
            ],
            [
                "Activity",
                "Actividad creada para una materia, con título, descripción, fecha de entrega y archivo de instrucciones.",
                "Permite mostrar tareas o trabajos pendientes asociados al estudiante o docente.",
            ],
            [
                "ActivitySubmission",
                "Entrega del estudiante, archivo, fecha, comentarios y calificación sobre 50.",
                "Permite saber si una actividad ya fue entregada y mostrar evidencia de entrega.",
            ],
            [
                "Notification",
                "Avisos dirigidos al estudiante, con estado leído/no leído y fecha.",
                "Alimenta indicadores y mensajes de alerta en la experiencia del dashboard.",
            ],
            [
                "ConversationHistory",
                "Historial de mensajes por estudiante, sesión, rol, contexto usado, tokens y fecha.",
                "Da continuidad conversacional a usuarios autenticados y permite trazabilidad del uso del asistente.",
            ],
        ],
        widths=[1.35, 2.7, 2.85],
    )

    add_heading(doc, "4.3 Endpoints académicos", 2)
    add_para(
        doc,
        "Los endpoints de estudiantes están protegidos por autenticación. Esto significa que la información académica personalizada solo está disponible cuando el usuario tiene un token válido y el backend puede validar su identidad.",
    )
    add_table(
        doc,
        ["Endpoint", "Función", "Resultado entregado"],
        [
            ["/api/students/profile/", "Consultar perfil del estudiante autenticado.", "Nombre, correo, carrera, semestre, código estudiantil y datos de perfil."],
            ["/api/students/grades/", "Consultar calificaciones por periodo académico.", "Materia, código, parciales, examen final, nota total, asistencia y estado."],
            ["/api/students/schedule/", "Consultar horario por periodo.", "Día, hora, materia, docente, aula, edificio y piso."],
            ["/api/students/stats/", "Calcular estadísticas académicas.", "Promedio, materias aprobadas, asistencia promedio, créditos y semestre actual."],
            ["/api/students/notifications/", "Listar notificaciones personales.", "Avisos del estudiante y estado de lectura."],
            ["/api/students/notifications/<id>/read/", "Marcar una notificación como leída.", "Actualización del estado de lectura."],
            ["/api/students/activities/", "Listar actividades de materias inscritas.", "Actividades pendientes o entregadas, fechas y estado de entrega."],
            ["/api/students/activities/<id>/submit/", "Enviar una actividad.", "Archivo de entrega, comentarios y registro de envío."],
            ["/api/students/teacher/subjects/", "Consultar materias del docente autenticado.", "Materias asignadas al profesor."],
            ["/api/students/teacher/schedule/", "Consultar horario docente.", "Materias, aulas y franjas horarias del docente."],
            ["/api/students/teacher/activities/", "Crear o listar actividades docentes.", "Gestión de tareas para materias asignadas."],
        ],
        widths=[2.2, 2.25, 2.45],
    )

    add_heading(doc, "4.4 Integración con el chatbot", 2)
    add_para(
        doc,
        "El servicio conversacional incorpora contexto académico mediante una función de construcción de contexto del estudiante. Cuando la solicitud llega con autenticación válida, el sistema agrega datos del estudiante, calificaciones, horario vigente, actividades pendientes y entregas. "
        "Con esa información, el modelo puede contestar de forma contextual y no repetir instrucciones genéricas.",
    )
    add_bullets(
        doc,
        [
            "Si el usuario pregunta por sus notas, el chatbot usa los registros Grade vinculados al estudiante autenticado.",
            "Si pregunta por horarios, se revisan los Schedule y Enrollment del periodo correspondiente.",
            "Si pregunta por actividades, se consultan Activity y ActivitySubmission para distinguir tareas pendientes y entregadas.",
            "Si la consulta es académica privada, el sistema requiere autenticación; en modo público no se exponen datos personales.",
            "El historial conversacional se almacena únicamente para estudiantes autenticados, evitando guardar historial sensible de usuarios anónimos.",
        ],
    )

    add_heading(doc, "4.5 Dashboard académico del estudiante", 2)
    add_para(
        doc,
        "En frontend, el sprint se refleja especialmente en el Dashboard del estudiante. Este dashboard organiza la información académica en una experiencia visual donde el usuario puede navegar entre materias, calendario, archivos y actividades. "
        "Además, integra el ChatWindow para que el estudiante consulte por lenguaje natural sin salir de su entorno académico.",
    )
    add_bullets(
        doc,
        [
            "El panel muestra materias y tarjetas de curso construidas desde las calificaciones y datos académicos del usuario.",
            "La vista de calendario organiza actividades y clases por días de la semana.",
            "La vista de actividad permite revisar detalles, descargar o consultar instrucciones y subir entregas.",
            "Las notificaciones no leídas alimentan indicadores visuales y pueden activar mensajes proactivos en el chat.",
            "El dashboard reutiliza la sesión JWT guardada por el frontend para acceder a endpoints protegidos.",
        ],
    )

    add_heading(doc, "4.6 Resultado del sprint", 2)
    add_para(
        doc,
        "Al finalizar este sprint, el proyecto cuenta con una capa académica funcional y consultable. El chatbot deja de ser únicamente informativo y se convierte en un asistente personalizado para estudiantes autenticados. "
        "La información personal se obtiene desde modelos estructurados y se entrega por endpoints protegidos, lo que permite mantener separación entre consultas públicas e información privada.",
    )


def add_sprint_2(doc: Document) -> None:
    add_heading(doc, "5. Sprint 2: Consultas institucionales", 1)
    add_para(
        doc,
        "Este sprint incorpora la arquitectura RAG para responder preguntas institucionales de PUCE-SI. El propósito es que el asistente no invente información cuando el usuario pregunte por servicios, ubicaciones, horarios institucionales, prerrequisitos, procesos académicos o normativa general. "
        "En lugar de depender solo del conocimiento del modelo, el sistema recupera fragmentos desde documentos validados y los incluye como contexto para generar la respuesta.",
    )

    add_heading(doc, "5.1 Objetivo funcional", 2)
    add_para(
        doc,
        "El objetivo es responder consultas públicas e institucionales con base documental. Esta capa es útil para usuarios no autenticados y autenticados, porque cubre información general de la universidad: servicios del campus, políticas de ingreso, gestión de carnet, horarios estudiantiles oficiales, prerrequisitos y datos institucionales.",
    )

    add_heading(doc, "5.2 Fuentes documentales del proyecto", 2)
    add_table(
        doc,
        ["Fuente", "Contenido", "Uso dentro del RAG"],
        [
            [
                "informacion_institucional.txt",
                "Información institucional base de PUCE-SI.",
                "Respalda respuestas generales sobre la sede y sus servicios.",
            ],
            [
                "puce_ibarra_info.txt",
                "Información complementaria sobre la sede Ibarra y datos institucionales.",
                "Permite responder preguntas frecuentes sobre la universidad y su contexto.",
            ],
            [
                "campus_mapa_servicios.txt",
                "Mapa textual de edificios, servicios, ingreso físico, carnet, espacios y reglas académicas de calificación.",
                "Responde consultas de ubicación, servicios, ingreso al campus, gestión de carnet y criterios de aprobación.",
            ],
            [
                "horarios_prerrequisitos_2026_01.txt",
                "Referencia oficial a horarios estudiantiles y prerrequisitos del periodo marzo-julio 2026.",
                "Guía al usuario hacia horarios oficiales públicos cuando la consulta no corresponde al horario personal autenticado.",
            ],
            [
                "scraped_web_content.txt",
                "Contenido obtenido desde páginas institucionales validadas mediante scraping controlado.",
                "Amplía el corpus RAG con información publicada por la universidad.",
            ],
        ],
        widths=[2.0, 2.7, 2.2],
    )

    add_heading(doc, "5.3 Información institucional añadida y corregida", 2)
    add_para(
        doc,
        "Dentro del corpus institucional se incorporaron datos específicos que son importantes para la operación real del campus. La redacción se mantiene profesional y orientada al estudiante.",
    )
    add_bullets(
        doc,
        [
            "La gestión del carnet estudiantil se ubica en el Edificio 2, piso 3.",
            "Para ingresar físicamente a la universidad se debe presentar el carnet estudiantil. En caso de no portarlo, el visitante o estudiante debe registrarse en la hoja de ingreso administrada por el personal de seguridad.",
            "La biblioteca se ubica en el Edificio 4.",
            "Los servicios de bar/cafetería se encuentran en los Edificios 1 y 4.",
            "El área de copias y laboratorios se ubica en el Edificio 3.",
            "Tesorería, secretaría y sala de grados se encuentran en el Edificio 4.",
            "Dirección de Estudiantes está en el Edificio 1.",
            "El estadio se encuentra hacia el Edificio 5 y la capilla después del ingreso principal.",
            "El sistema de calificación se documenta sobre 200 puntos, con cuatro componentes de 50 puntos, mínimo de 30/50 por componente, mínimo total de 120/200 y asistencia mínima de 80%.",
        ],
    )

    add_heading(doc, "5.4 Funcionamiento técnico del RAG", 2)
    add_para(
        doc,
        "El motor RAG se implementa con indexación de fragmentos documentales, embeddings de OpenAI y búsqueda vectorial con FAISS. Cuando se construye el índice, los documentos se dividen en fragmentos, se generan representaciones semánticas y se almacenan para búsqueda posterior. "
        "Cuando el usuario pregunta, la consulta se transforma en embedding, se recuperan fragmentos relevantes y esos fragmentos se envían al prompt del chatbot como contexto verificable.",
    )
    add_numbered(
        doc,
        [
            "Carga de documentos desde la carpeta institucional del backend.",
            "Exclusión de archivos no aptos para RAG, como datos de prueba de estudiantes.",
            "Fragmentación textual para mejorar precisión de recuperación.",
            "Generación de embeddings mediante el proveedor configurado.",
            "Construcción o lectura del índice FAISS.",
            "Búsqueda por similitud entre consulta y documentos.",
            "Formateo de contexto con fuentes para el servicio conversacional.",
            "Generación de respuesta en español con instrucción de no inventar cuando no exista información suficiente.",
        ],
    )

    add_heading(doc, "5.5 Respaldo de búsqueda léxica", 2)
    add_para(
        doc,
        "El proyecto incorpora un respaldo de búsqueda léxica para evitar que el RAG falle completamente si no se puede generar un embedding o si el índice está desactualizado. "
        "Este respaldo normaliza texto, elimina diferencias por acentos, usa coincidencia de términos y expande consultas relacionadas con ingreso, acceso, carnet y servicios del campus. "
        "Con esto se mejora la robustez para preguntas concretas como dónde gestionar el carnet o cómo ingresar físicamente a la universidad.",
    )

    add_heading(doc, "5.6 Endpoints RAG", 2)
    add_table(
        doc,
        ["Endpoint", "Permiso", "Responsabilidad"],
        [
            ["/api/rag/build/", "Administrador", "Construir o reconstruir el índice RAG y registrar el evento de auditoría correspondiente."],
            ["/api/rag/search/", "Usuario autenticado", "Ejecutar una búsqueda RAG para inspección o depuración de resultados recuperados."],
        ],
        widths=[2.2, 1.7, 3.0],
    )

    add_heading(doc, "5.7 Relación con el prompt del chatbot", 2)
    add_para(
        doc,
        "El prompt de sistema indica que las respuestas deben darse en español, con tono claro y profesional. También establece que el asistente no debe inventar datos y que debe usar el contexto RAG cuando la consulta sea institucional. "
        "Además, se define una jerarquía de fuentes: primero documentos institucionales validados y datos estructurados del sistema; luego información del historial conversacional cuando aplica; y finalmente una negativa clara si la información no se encuentra.",
    )

    add_heading(doc, "5.8 Resultado del sprint", 2)
    add_para(
        doc,
        "Al finalizar este sprint, el chatbot puede responder consultas institucionales usando una base documental actualizable. Esto permite que el sistema tenga una fuente de verdad para temas públicos de PUCE-SI y reduce el riesgo de respuestas inventadas. "
        "La arquitectura queda preparada para añadir nuevos documentos oficiales o páginas institucionales conforme evolucione la universidad.",
    )


def add_sprint_3(doc: Document) -> None:
    add_heading(doc, "6. Sprint 3: Interacción conversacional por voz y respuesta hablada", 1)
    add_para(
        doc,
        "El título original del sprint era reconocimiento de voz. Sin embargo, el proyecto implementa más que reconocimiento: permite grabar audio, transcribirlo, procesarlo con el chatbot, generar respuesta textual y devolver audio hablado. "
        "Por eso el nombre más preciso para la tesis es Sprint 3: Interacción conversacional por voz y respuesta hablada.",
    )

    add_heading(doc, "6.1 Objetivo funcional", 2)
    add_para(
        doc,
        "El objetivo es hacer que el chatbot sea accesible mediante voz. El usuario puede mantener presionado el micrófono, grabar su consulta, enviarla al backend, recibir la transcripción, obtener la respuesta del asistente y escucharla en audio. "
        "Esto mejora la accesibilidad y acerca el sistema a una interacción más natural.",
    )

    add_heading(doc, "6.2 Servicio de transcripción", 2)
    add_para(
        doc,
        "El backend cuenta con un servicio de voz que usa reconocimiento automático de habla mediante OpenAI. El archivo de audio se valida, se envuelve como archivo webm y se procesa con configuración de idioma español. "
        "El resultado principal es texto limpio, listo para ser enviado al mismo servicio conversacional que atiende las consultas escritas.",
    )
    add_bullets(
        doc,
        [
            "Valida que el audio exista y tenga un tamaño mínimo razonable antes de enviarlo a transcripción.",
            "Usa el modelo de transcripción configurado para convertir audio en texto.",
            "Mantiene el idioma español para mejorar la precisión en nombres, servicios y consultas académicas.",
            "Devuelve errores controlados cuando el audio es inválido o no puede ser procesado.",
        ],
    )

    add_heading(doc, "6.3 Servicio de síntesis de voz", 2)
    add_para(
        doc,
        "Después de obtener la respuesta del chatbot, el servicio de voz genera audio en formato MP3. Antes de sintetizar, el texto se normaliza para mejorar la pronunciación de siglas y términos técnicos frecuentes dentro del proyecto, como PUCE-SI, PUCESI, PDF, NRC, OTP, 2FA, IA, URL y correo electrónico.",
    )
    add_bullets(
        doc,
        [
            "Usa un modelo de texto a voz configurado para español.",
            "Incluye instrucciones de voz clara, natural y profesional.",
            "Tiene mecanismo de respaldo con otro modelo de TTS si el modelo principal no está disponible.",
            "Devuelve bytes de audio que luego se codifican en base64 para el frontend.",
        ],
    )

    add_heading(doc, "6.4 Endpoints de voz", 2)
    add_table(
        doc,
        ["Endpoint", "Acceso", "Flujo"],
        [
            ["/api/voice/stt/", "Autenticado", "Recibe un archivo de audio, valida tamaño, transcribe la consulta y devuelve texto."],
            ["/api/voice/chat/", "Público o autenticado", "Recibe audio, transcribe, consulta el chatbot, genera audio de respuesta y devuelve transcript, texto y audio_base64."],
        ],
        widths=[2.1, 1.8, 3.0],
    )

    add_heading(doc, "6.5 Integración en frontend", 2)
    add_para(
        doc,
        "La interfaz conversacional incluye un botón de micrófono. El navegador solicita permiso para usar el micrófono y graba mediante MediaRecorder. El usuario puede mantener presionado para grabar y soltar para enviar. "
        "Luego el componente reproduce el audio devuelto por el backend, deteniendo cualquier reproducción anterior para evitar solapamientos.",
    )
    add_bullets(
        doc,
        [
            "La grabación usa APIs nativas del navegador: getUserMedia y MediaRecorder.",
            "El audio se envía como FormData al endpoint de voz.",
            "La reproducción usa WebAudio cuando es posible y un elemento Audio como respaldo.",
            "El componente maneja bloqueos del navegador cuando la reproducción automática requiere interacción del usuario.",
            "El sistema muestra avisos controlados si no hay permiso de micrófono o si la grabación no puede procesarse.",
        ],
    )

    add_heading(doc, "6.6 Resultado del sprint", 2)
    add_para(
        doc,
        "Al finalizar este sprint, el chatbot no depende exclusivamente del teclado. Puede recibir preguntas por voz, responder con texto y reproducir audio, manteniendo la misma lógica de autenticación, RAG y contexto académico. "
        "Esto convierte la voz en una extensión de la arquitectura conversacional, no en un módulo aislado.",
    )


def add_sprint_4(doc: Document) -> None:
    add_heading(doc, "7. Sprint 4: Interfaz conversacional", 1)
    add_para(
        doc,
        "Este sprint reúne la experiencia visual y conversacional del sistema. Incluye la ventana de chat, la navegación pública, el inicio de sesión, el dashboard de estudiante, el panel docente y la integración con el cliente API. "
        "La meta es que el usuario pueda interactuar con el chatbot de forma natural y que las respuestas se conecten con el contexto correcto según el estado de autenticación.",
    )

    add_heading(doc, "7.1 Objetivo funcional", 2)
    add_para(
        doc,
        "El objetivo es ofrecer una interfaz clara para consultas públicas y privadas. Un usuario externo puede consultar información institucional desde páginas públicas, mientras que un estudiante o docente autenticado puede acceder a vistas protegidas y a respuestas más contextualizadas.",
    )

    add_heading(doc, "7.2 Estructura de rutas", 2)
    add_table(
        doc,
        ["Ruta", "Tipo", "Función"],
        [
            ["/", "Pública", "Página principal con acceso al chatbot público e información general."],
            ["/login", "Pública", "Inicio de sesión con credenciales y verificación OTP cuando corresponde."],
            ["/banner", "Pública", "Página informativa asociada a autoservicio o comunicación institucional."],
            ["/empleo", "Pública", "Página de bolsa de empleo o información relacionada."],
            ["/dashboard", "Protegida", "Vista académica para estudiantes y administradores autorizados."],
            ["/teacher", "Protegida", "Panel docente para consultar materias, horarios y gestionar actividades."],
        ],
        widths=[1.35, 1.35, 4.2],
    )

    add_heading(doc, "7.3 Ventana de chat", 2)
    add_para(
        doc,
        "El componente ChatWindow es el núcleo de la interfaz conversacional. Funciona como una ventana flotante con encabezado institucional, estado público o de estudiante, listado de mensajes, indicador de carga, campo de texto, botón de envío y botón de micrófono. "
        "Este componente puede usarse en modo público o autenticado, lo que permite reutilizar la experiencia de chat en varias páginas.",
    )
    add_bullets(
        doc,
        [
            "Mantiene estado local de mensajes, entrada de texto, carga, grabación, sesión conversacional y avisos de audio.",
            "Envía preguntas escritas al endpoint de chatbot.",
            "Envía preguntas habladas al endpoint de voz.",
            "Renderiza burbujas diferenciadas para usuario y asistente.",
            "Permite mensajes iniciales y notificaciones proactivas cuando el contexto lo requiere.",
            "Gestiona reproducción de audio para respuestas generadas por TTS.",
        ],
    )

    add_heading(doc, "7.4 Cliente API y sesión", 2)
    add_para(
        doc,
        "El frontend usa Axios como cliente HTTP. Antes de cada solicitud, el interceptor agrega el access token JWT cuando existe. Si una solicitud falla con 401, el interceptor intenta refrescar la sesión usando el refresh token. "
        "Si el refresco falla, limpia el almacenamiento local y redirige al login. Este comportamiento evita que el usuario pierda sesión innecesariamente y mantiene la protección de endpoints privados.",
    )
    add_bullets(
        doc,
        [
            "El estado de autenticación se gestiona con Zustand.",
            "Los tokens access_token y refresh_token se guardan en localStorage.",
            "El método setAuth almacena usuario y tokens después de login u OTP.",
            "El método logout elimina tokens y reinicia el estado de sesión.",
            "Las rutas protegidas validan autenticación y rol antes de mostrar contenido privado.",
        ],
    )

    add_heading(doc, "7.5 Login y verificación", 2)
    add_para(
        doc,
        "La página de login contempla dos etapas. Primero solicita credenciales. Si el usuario tiene 2FA activo, el backend devuelve un temp_token y el frontend muestra la segunda etapa para ingresar el código OTP. "
        "Después de verificar el OTP, el sistema recibe tokens JWT, guarda la sesión y redirige según rol: estudiante o administrador hacia dashboard, docente hacia el panel docente.",
    )

    add_heading(doc, "7.6 Dashboard estudiantil", 2)
    add_para(
        doc,
        "El dashboard estudiantil centraliza la experiencia académica. No es solo una pantalla de datos; está diseñado para que el usuario pueda revisar cursos, navegar por calendario, consultar archivos, revisar actividades y comunicarse con el chatbot desde el mismo entorno.",
    )
    add_bullets(
        doc,
        [
            "Presenta materias derivadas de los datos académicos del estudiante.",
            "Permite seleccionar cursos y revisar detalles asociados.",
            "Muestra próximas actividades, eventos y líneas de tiempo.",
            "Incluye carga de archivos para entregas de actividades.",
            "Integra el chat como apoyo contextual dentro de la vida académica del estudiante.",
        ],
    )

    add_heading(doc, "7.7 Panel docente", 2)
    add_para(
        doc,
        "El panel docente permite que un profesor autenticado consulte sus materias, vea horarios, revise actividades y cree nuevas tareas. "
        "Esto amplía el sistema más allá del estudiante y permite administrar información académica que luego puede ser consultada por quienes estén inscritos en las materias correspondientes.",
    )
    add_bullets(
        doc,
        [
            "Lista materias asignadas al docente.",
            "Consulta horarios y aulas.",
            "Permite crear actividades con título, descripción, fecha de entrega y archivo opcional.",
            "Consulta actividades existentes y sus datos principales.",
        ],
    )

    add_heading(doc, "7.8 Respuestas de fecha, hora y clima", 2)
    add_para(
        doc,
        "Como parte de la experiencia conversacional, el chatbot incorpora herramientas internas para responder preguntas de fecha, hora y clima. "
        "Esto corrige el comportamiento en el que el asistente respondía que no tenía acceso a la fecha actual. El sistema ahora puede responder directamente a preguntas como qué día es hoy, qué día fue ayer, qué hora es y cuál es el clima de una ciudad soportada.",
    )
    add_bullets(
        doc,
        [
            "La zona horaria base del proyecto es America/Guayaquil.",
            "El contexto del sistema incluye fecha actual, hora actual, día anterior y día siguiente.",
            "Las preguntas directas de fecha y hora pueden responderse sin depender del modelo de lenguaje.",
            "La consulta de clima usa Open-Meteo para ciudades soportadas, con Ibarra como ubicación predeterminada.",
            "Las preguntas académicas como clases de hoy no se interceptan como simple fecha; se dejan pasar al contexto académico cuando el usuario está autenticado.",
        ],
    )

    add_heading(doc, "7.9 Resultado del sprint", 2)
    add_para(
        doc,
        "Al finalizar este sprint, el usuario cuenta con una interfaz coherente para hablar con el asistente, iniciar sesión, acceder a datos académicos y usar la voz. "
        "La capa visual se integra con seguridad, RAG, datos académicos y servicios de IA, por lo que la experiencia no queda fragmentada entre pantallas aisladas.",
    )


def add_sprint_5(doc: Document) -> None:
    add_heading(doc, "8. Sprint 5: Seguridad y auditoría", 1)
    add_para(
        doc,
        "Este sprint protege el sistema y deja evidencia de uso. Es especialmente importante porque el proyecto maneja información académica personal, credenciales, archivos de entrega, historial conversacional y servicios de IA que podrían ser abusados si no tuvieran controles.",
    )

    add_heading(doc, "8.1 Objetivo funcional", 2)
    add_para(
        doc,
        "El objetivo es implementar una seguridad completa y funcional para el prototipo de tesis: autenticación robusta, autorización por roles, protección de contraseñas y OTP, bloqueo ante intentos fallidos, límites de uso, controles de archivos, configuración segura por ambiente y auditoría de acciones relevantes.",
    )

    add_heading(doc, "8.2 Autenticación y roles", 2)
    add_para(
        doc,
        "El modelo de usuario extiende el usuario base de Django e incorpora UUID como identificador, correo único, rol, estado de 2FA, secreto OTP, IP de último login, contador de intentos fallidos y fecha de bloqueo. "
        "Los roles principales son estudiante, docente y administrador. El sistema usa estos roles para decidir qué rutas y endpoints puede utilizar cada usuario.",
    )
    add_bullets(
        doc,
        [
            "El estudiante accede a consultas académicas propias, dashboard y chat personalizado.",
            "El docente accede al panel docente, materias, horarios y actividades.",
            "El administrador puede acceder a operaciones sensibles como reconstrucción RAG y consulta de auditoría.",
            "La propiedad effective_role permite interpretar docentes aunque su rol base no esté sincronizado, siempre que exista relación con Professor.",
        ],
    )

    add_heading(doc, "8.3 JWT y ciclo de sesión", 2)
    add_para(
        doc,
        "La autenticación usa JWT. El access token tiene una duración corta y el refresh token una duración mayor. El refresh se rota y el token anterior se incluye en blacklist cuando corresponde. "
        "Esta estrategia reduce exposición si un token de acceso se filtra y permite mantener sesiones de usuario sin pedir credenciales constantemente.",
    )
    add_bullets(
        doc,
        [
            "Access token configurado con duración de 15 minutos.",
            "Refresh token configurado con duración de 7 días.",
            "Rotación de refresh token habilitada.",
            "Blacklist de refresh token después de rotación.",
            "Actualización del último login cuando se emiten tokens.",
        ],
    )

    add_heading(doc, "8.4 Contraseñas, OTP y bloqueo", 2)
    add_para(
        doc,
        "Las contraseñas se protegen con Argon2 como algoritmo principal de hashing y validadores de contraseña de Django. "
        "Para usuarios con doble factor de autenticación, el sistema genera códigos OTP con aleatoriedad segura. El código se almacena hasheado y no como valor plano, lo que evita que un acceso directo a la base de datos revele el OTP activo.",
    )
    add_bullets(
        doc,
        [
            "Después de cinco intentos fallidos, la cuenta se bloquea temporalmente.",
            "El bloqueo tiene una duración de 15 minutos.",
            "El OTP se marca como usado después de validarse.",
            "El OTP tiene fecha de expiración.",
            "El temp_token evita que la verificación OTP dependa solo de usuario y contraseña.",
            "En desarrollo, el envío de correos puede usarse con backend de consola para pruebas controladas.",
        ],
    )

    add_heading(doc, "8.5 Protección de endpoints y límites de abuso", 2)
    add_para(
        doc,
        "El backend usa throttling de Django REST Framework para limitar el abuso de endpoints. Esto es importante en login, chat y voz, porque esas rutas podrían ser atacadas por fuerza bruta o generar costos innecesarios de servicios de IA.",
    )
    add_table(
        doc,
        ["Control", "Configuración", "Riesgo mitigado"],
        [
            ["AnonRateThrottle", "20 solicitudes por minuto", "Abuso desde usuarios no autenticados."],
            ["UserRateThrottle", "100 solicitudes por minuto", "Uso excesivo de usuarios autenticados."],
            ["login", "5 solicitudes por minuto", "Fuerza bruta de credenciales."],
            ["chat", "10 solicitudes por minuto", "Abuso del modelo conversacional y costos de IA."],
            ["voice", "6 solicitudes por minuto", "Abuso de transcripción y síntesis de voz."],
        ],
        widths=[1.8, 2.0, 3.1],
    )

    add_heading(doc, "8.6 Protección de datos y archivos", 2)
    add_para(
        doc,
        "La protección de datos incluye validación de longitud de mensajes, límites de carga de audio, límites de archivos de actividades y extensiones permitidas. "
        "Estos controles evitan entradas excesivas, archivos no esperados y consumo innecesario de memoria o procesamiento.",
    )
    add_bullets(
        doc,
        [
            "El chatbot valida longitud máxima del mensaje antes de procesarlo.",
            "El servicio de voz valida tamaño máximo de audio.",
            "La entrega de actividades valida tamaño y extensiones permitidas.",
            "Django limita tamaño de subida en memoria para reducir riesgo de agotamiento de recursos.",
            "La información académica privada no se entrega en modo público.",
        ],
    )

    add_heading(doc, "8.7 Configuración segura por ambiente", 2)
    add_para(
        doc,
        "El archivo de configuración separa valores sensibles mediante variables de entorno. En producción se pueden activar cookies seguras, CSRF seguro, redirección SSL, HSTS, bloqueo de framing, protección contra sniffing de contenido y políticas de origen. "
        "El entorno de desarrollo mantiene flexibilidad, pero permite que DJANGO_DEBUG sea configurable y no quede permanentemente activado.",
    )
    add_bullets(
        doc,
        [
            "CORS se controla mediante orígenes permitidos configurables.",
            "CSRF_TRUSTED_ORIGINS se configura desde variables de entorno.",
            "SECURE_SSL_REDIRECT, SESSION_COOKIE_SECURE y CSRF_COOKIE_SECURE pueden habilitarse para despliegue.",
            "X_FRAME_OPTIONS se establece en DENY para evitar clickjacking.",
            "SECURE_CONTENT_TYPE_NOSNIFF reduce riesgos por interpretación incorrecta de contenido.",
            "REFERRER_POLICY y CROSS_ORIGIN_OPENER_POLICY refuerzan aislamiento del navegador.",
        ],
    )

    add_heading(doc, "8.8 Auditoría", 2)
    add_para(
        doc,
        "El sistema registra eventos relevantes en AuditLog. Cada registro incluye usuario cuando existe, acción, severidad, IP, user agent, metadatos y fecha. "
        "La auditoría ayuda a revisar accesos, intentos fallidos, consultas al chatbot, reconstrucción de RAG y accesos a datos sensibles.",
    )
    add_table(
        doc,
        ["Acción auditada", "Descripción", "Importancia"],
        [
            ["LOGIN_SUCCESS", "Inicio de sesión exitoso.", "Permite rastrear accesos válidos."],
            ["LOGIN_FAILED", "Intento fallido de login.", "Ayuda a detectar fuerza bruta o errores recurrentes."],
            ["LOGIN_OTP_SENT", "Envío de código OTP.", "Permite verificar activación del segundo factor."],
            ["LOGOUT", "Cierre de sesión.", "Registra finalización de sesión."],
            ["ACCOUNT_LOCKED", "Bloqueo temporal de cuenta.", "Evidencia intentos repetidos fallidos."],
            ["CHAT_QUERY", "Consulta textual al chatbot.", "Mide uso y permite análisis de interacciones."],
            ["CHAT_VOICE_QUERY", "Consulta por voz.", "Registra uso de la modalidad hablada."],
            ["DATA_ACCESS_GRADES", "Acceso a calificaciones.", "Protege trazabilidad de datos académicos sensibles."],
            ["DATA_ACCESS_SCHEDULE", "Acceso a horarios.", "Permite auditar información académica personal."],
            ["DATA_ACCESS_PROFILE", "Acceso al perfil.", "Registra consulta de información personal."],
            ["RAG_QUERY", "Consulta al RAG.", "Permite analizar recuperación institucional."],
            ["RAG_INDEX_BUILD", "Construcción del índice RAG.", "Evidencia cambios sobre la base documental."],
            ["SYSTEM_ERROR", "Error del sistema.", "Soporta análisis técnico y diagnóstico."],
        ],
        widths=[1.8, 2.6, 2.5],
    )

    add_heading(doc, "8.9 Endpoints administrativos de auditoría", 2)
    add_table(
        doc,
        ["Endpoint", "Permiso", "Responsabilidad"],
        [
            ["/api/logs/", "Administrador", "Listar logs con paginación y filtros por acción o severidad."],
            ["/api/logs/chatbot-usage/", "Administrador", "Consultar estadísticas de uso del chatbot filtradas por usuario o rango de fechas."],
        ],
        widths=[2.2, 1.7, 3.0],
    )

    add_heading(doc, "8.10 Resultado del sprint", 2)
    add_para(
        doc,
        "Al finalizar este sprint, el proyecto dispone de una base de seguridad funcional para una tesis: identidad protegida, segundo factor, tokens rotativos, bloqueo por intentos, límites por endpoint, protección de archivos y trazabilidad administrativa. "
        "Estos controles no reemplazan una auditoría profesional de producción, pero sí elevan el prototipo por encima de una prueba básica sin seguridad real.",
    )


def add_traceability(doc: Document) -> None:
    add_heading(doc, "9. Trazabilidad entre sprints y módulos", 1)
    add_para(
        doc,
        "La siguiente tabla resume qué partes del proyecto sostienen cada sprint. Esta trazabilidad permite defender en la tesis que las funcionalidades descritas están implementadas en componentes concretos.",
    )
    add_table(
        doc,
        ["Sprint", "Módulos y archivos representativos", "Evidencia funcional"],
        [
            [
                "Sprint 1",
                "backend/apps/students, Student, Grade, Schedule, Activity, ActivitySubmission, Notification, ConversationHistory, DashboardPage, TeacherDashboardPage.",
                "Consultas de perfil, notas, horarios, actividades, entregas, estadísticas y contexto académico en chatbot autenticado.",
            ],
            [
                "Sprint 2",
                "backend/apps/rag, backend/documents, campus_mapa_servicios.txt, horarios_prerrequisitos_2026_01.txt, scraped_web_content.txt.",
                "Recuperación de información institucional validada, gestión de carnet, ingreso físico al campus, servicios y horarios oficiales.",
            ],
            [
                "Sprint 3",
                "backend/apps/voice, voice/services.py, voice/views.py, ChatWindow con MediaRecorder y reproducción de audio.",
                "Audio a texto, consulta conversacional y respuesta hablada.",
            ],
            [
                "Sprint 4",
                "src/App.tsx, src/components/ChatWindow.tsx, src/pages/LandingPage.tsx, LoginPage, DashboardPage, TeacherDashboardPage, api.ts, authStore.ts.",
                "Interfaz pública y privada, chat flotante, login con OTP, rutas protegidas, dashboard y panel docente.",
            ],
            [
                "Sprint 5",
                "backend/apps/authentication, backend/apps/audit_logs, config/settings, serializers, services, throttling, JWT, OTP, logs.",
                "Seguridad de sesión, roles, OTP hasheado, bloqueo, rate limits, auditoría y endpoints administrativos.",
            ],
        ],
        widths=[0.9, 3.15, 2.85],
    )

    add_heading(doc, "10. Validaciones realizadas en el proyecto", 1)
    add_para(
        doc,
        "Durante el análisis y corrección del proyecto se verificaron elementos clave de funcionamiento. El frontend compila correctamente y el linter no reporta errores bloqueantes. También se comprobaron respuestas directas de fecha y clima, incluyendo respuestas para hoy, ayer y clima de Ibarra mediante Open-Meteo. "
        "Además, se revisó que el RAG cuente con respaldo léxico para consultas institucionales específicas y que la información de carnet e ingreso físico al campus esté incluida en documentos validados.",
    )
    add_bullets(
        doc,
        [
            "Compilación frontend completada correctamente.",
            "Lint frontend sin errores; permanecen advertencias no bloqueantes relacionadas con dependencias de hooks.",
            "Respuestas directas de fecha y hora operan con zona horaria America/Guayaquil.",
            "Consulta de clima de Ibarra funciona mediante servicio externo Open-Meteo cuando hay red disponible.",
            "La búsqueda RAG tiene respaldo léxico para preguntas de carnet, ingreso, acceso y servicios del campus.",
            "La seguridad incluye OTP hasheado, throttling por alcance, bloqueo temporal y auditoría.",
        ],
    )

    add_heading(doc, "11. Conclusión", 1)
    add_para(
        doc,
        "Los cinco sprints conforman una evolución coherente del prototipo. Primero se crea la base académica personalizada; luego se añade información institucional mediante RAG; después se amplía la interacción con voz; posteriormente se consolida la interfaz conversacional; y finalmente se fortalecen seguridad y auditoría. "
        "El resultado es un chatbot académico que puede funcionar como asistente público para consultas institucionales y como asistente autenticado para información académica personalizada, manteniendo controles de protección de datos y trazabilidad.",
    )


def build_document() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_intro(doc)
    add_sprint_1(doc)
    add_sprint_2(doc)
    add_sprint_3(doc)
    add_sprint_4(doc)
    add_sprint_5(doc)
    add_traceability(doc)

    for section in doc.sections:
        section.different_first_page_header_footer = False
        section.footer_distance = Cm(1.1)

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
